import streamlit as st # frontend library

from streamlit_chat import message #helps to display messages on frontend or app

import os #to interect with operating system

from PyPDF2  import PdfReader #to working with pdf files

import docx #to working with Microsoft Word Docs

from langchain.chat_models import ChatOpenAI #to connect our chat bot to llm(GPT)

from langchain.llms import OpenAI #just another GPT model

from dotenv import load_dotenv #hepls to load env variables

from langchain.embeddings import HuggingFaceEmbeddings #to convert txt into numerical vecotrs(representations)

from langchain import HuggingFaceHub #to connect huggingface pre-trained models and datasets(resources)

from langchain.text_splitter import CharacterTextSplitter #to break down large peace of text into smaller chunks

from langchain.vectorstores import FAISS #Facebook AI Similarity Search tool to store and search vectors efficiently.

from langchain.chains import ConversationalRetrievalChain #to combine language model,vector store,memory  relevent info during conversation.
from langchain.chains import RetrievalQA

from langchain.memory import ConversationBufferMemory #track previous conversation and allowing the chatbot what user already said .

from langchain.callbacks import get_openai_callback

from sentence_transformers import SentenceTransformer


open_api_key = st.secrets['OPENAI_API_KEY']

def main():
    load_dotenv()
    st.set_page_config(page_title = 'Document GPT')
    st.title('Your Document GPT')

    if 'conversation' not in st.session_state:
        st.session_state.conversation = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = None
    if 'processComplete' not in st.session_state:
        st.session_state.processComplete = None
    
    with st.sidebar:
        uploaded_files = st.file_uploader('Upload Your File here.',type = ['pdf'],accept_multiple_files = True)
        openai_api_key = open_api_key
        # openai_api_key = st.text_input('Your Open AI Api Key',key = open_api_key,type = 'password')

        process = st.button('Run')
    if process:
        if not openai_api_key:
            st.error('Please provide API key first.')
            st.stop()
        files_text = get_files_text(uploaded_files)
        st.write('File Loadeding...')
        # get text chunks
        text_chunks = get_text_chunks(files_text)
        st.write('Files Chunks created...')
        # create vector store 
        vectorstore = get_vector_store(text_chunks)
        st.write('vectore store created...')

        # create conversation chain 
        st.session_state.conversation = get_conversation_chain(vectorstore,open_api_key)

        st.session_state.processComplete = True

    if st.session_state.processComplete == True:
        user_question = st.chat_input('Ask Questions about your files.')
        if user_question:
            handle_userinput(user_question)

# functions

def get_files_text(uploaded_files):
    # set text empty initially
    text = ""
    for uploaded_file in uploaded_files:

        splited_name = os.path.splitext(uploaded_file.name) #it split the file into 'name' and 'extension'

        file_extension = splited_name[1] # [0] for 'file name' and [1] for 'file extension'

        if file_extension == '.pdf':
            text = text + get_pdf_text(uploaded_file)

        elif file_extension == '.docx':
            text = text + get_docx_text(uploaded_file)

        elif file_extension == '.csv':
            text = text + get_csv_text(uploaded_file)
        
        else:
            st.error('Please provide file with extension pdf,docx and csv')
    
    return text

# functions to read files children functions of get_files_text()

def get_pdf_text(file_extension):
    pdf_reader = PdfReader(file_extension)
    text = ""
    for page in pdf_reader.pages: #the .pages attribute contains all the pages of pdf
        text = text + page.extract_text()
    return text 

def get_docx_text(file_extension):
    doc = docx.Document(file_extension) #loading the docx file using python_docx library to get the docx file text
    allText = []
    for docparagraph in doc.paragraphs: #.paragraphs attribute contains all the paragraphs of docx file
        allText.append(docparagraph.text)
    text = ' '.join(allText)
    return text

def get_csv_text(file_extension):
    return 'Khud krain'

def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator = '\n',
        chunk_size = 800,
        chunk_overlap = 110,
        length_function = len
    )
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    # using hugging face model 
    embeddings = HuggingFaceEmbeddings(model_name = 'all-MiniLM-L6-v2')
    # create vector store using facebook's FAISS store
    knowledge_base = FAISS.from_texts(text_chunks,embeddings)
    return knowledge_base

# def get_conversation_chain(vectorstore,open_api_key):
#     # create an llm
#     created_llm = ChatOpenAI(openai_api_key = open_api_key,model_name = 'gpt-3.5-turbo',temperature = 0)
#     #create memory to store the past conversation history|storing history in chat history
#     created_memory = ConversationBufferMemory(memory_key = 'chat_history',return_messages = True)# return true just to return the messaege
#     # now we'r going to talk to vector store
#     conversation_chain = RetrievalQA(
#         llm = created_llm,
#         retriever = vectorstore.as_retriever(),#this is for ,to open the door of our vectorstore
#         chain_type = 'stuff'
#         # memory = created_memory,
#         # combine_docs_chain = None,
#         # question_generator = None
#     )
#     return conversation_chain

from langchain.chains import LLMChain, StuffDocumentsChain
from langchain.prompts import PromptTemplate
from langchain.chains import ConversationalRetrievalChain

def get_conversation_chain(vectorstore, open_api_key):
    # Create an LLM
    created_llm = ChatOpenAI(openai_api_key=open_api_key, model_name='gpt-3.5-turbo', temperature=0)

    # Create a prompt template for the question generator
    question_template = PromptTemplate(template="Question: {question}", input_variables=["question"])
    question_generator = LLMChain(llm=created_llm, prompt=question_template)

    # Create a prompt template for combining documents
    combine_template = PromptTemplate(template="Combine these documents: {documents}", input_variables=["documents"])
    combine_llm_chain = LLMChain(llm=created_llm, prompt=combine_template)

    # Initialize the StuffDocumentsChain
    combine_docs_chain = StuffDocumentsChain(llm_chain=combine_llm_chain)

    # Initialize the ConversationalRetrievalChain
    conversation_chain = ConversationalRetrievalChain(
        retriever=vectorstore.as_retriever(),
        memory=ConversationBufferMemory(memory_key='chat_history', return_messages=True),
        combine_docs_chain=combine_docs_chain,
        question_generator=question_generator
    )
    return conversation_chain


def handle_userinput(user_question):
    with get_openai_callback() as cb: #it is used to track the cost and tokens used during OpenAI API call.
        # 'with' is used to open callback and automatically close the when it done.
        response = st.session_state.conversation ({'question':user_question}) #this access the conversation chain and store all conversation in the respnse
        st.session_state.chat_history = response['chat_history']

        # layout of input 
        response_box = st.container()
        with response_box:
            for i,messages in enumerate(st.session_state.chat_history):
                if i % 2 == 0:
                    message(messages.content,us_user = True,key = str(i))
                else:
                    message(messages.content,key = str(i))

if __name__ == '__main__':
    main()








