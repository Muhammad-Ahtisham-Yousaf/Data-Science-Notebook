import streamlit as st

import os #used to interect with opreting system

from PyPDF2 import PdfReader#used to read,extract and manipulate pdf files

import docx #used to read,extract and work with microsoft word

from langchain.chat_models import ChatOpenAI #to use llm(gpt)

from langchain.vectorstores import FAISS #Qdrant

import faiss

# from qdrant_client import QdrantClient #Qdrant is a vectore database used to store and search large amounts of data using embeddings

import random #used to generate random numbers

from datetime import datetime #to work with datas and time

from langchain import PromptTemplate #to make sure model gets structured prompts

import string # provide utilities for working with strings

from dotenv import load_dotenv #for loading senstive info like API key securely using .env file

from langchain.embeddings import HuggingFaceEmbeddings #hugging face model used for embeddings( to convert text into numaric form ) to store in database like Qdrant

from langchain.text_splitter import CharacterTextSplitter#splits large text into smaller chunks for better processing

from streamlit_chat import message # for displaying chat-like text in streamlit app (frontend)

from langchain.docstore.document import Document#organizes metadata like titles or tags,so it's easier to search

import numpy as np



my_openai_key = st.secrets['OPENAI_API_KEY']
# qdrant_url = st.secrets['QDRAND_URL']
# qdrand_api_key = st.secrets['QDRANT_API_KEY']

def main():
    load_dotenv()
    st.set_page_config('Q/A ChatBot')
    st.title('Q/A ChatBot')

    if 'conversation' not in st.session_state:
        st.session_state.conversation = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'processComplete' not in st.session_state:
        st.session_state.processComplete = None

    with st.sidebar:
        uploaded_files = st.file_uploader('Uploade Your Files here:',key = ['pdf'], accept_multiple_files=True)
    
        openai_api_key = my_openai_key
        process = st.button('Start Process')
        # st.write(uploaded_files)
        
        if process:
            if uploaded_files == []:
                st.error('Please Upload file first,Thanks!')
                st.stop()
            if openai_api_key == "":
                st.error('Please provide Your Open API key first')
                st.stop()
        
            text_chunks_list = []
            for uploaded_file in uploaded_files:
                file_name = uploaded_file.name
                file_text = get_files_text(uploaded_file)
                
                text_chunks = get_text_chunks(file_text,file_name)
                text_chunks_list.extend(text_chunks)

                # st.write(text_chunks_list)
                curr_date = str(datetime.now())
                collection_name = "".join(random.choices(string.ascii_letters , k=4 )) + curr_date.split('.')[0].replace(':','-').replace(" ",'T')
                # st.write(collection_name) 
                # st.write(curr_date)
                # st.write(text_chunks_list) 
                vector_store = get_vector_store(text_chunks_list)
                st.write('Creating Vectors')

def get_files_text(uploaded_file):
    text = ""
    splite_tup = os.path.splitext(uploaded_file.name)
    # st.write(text)
    # st.subheader(splite_tup[0])
    file_extension = splite_tup[1]
    if file_extension == '.pdf':
        text += get_pdf_text(uploaded_file)
    if file_extension == '.docx':
        text += get_docx_text(uploaded_file)
    return text


def get_pdf_text(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    
    return text

def get_docx_text(docx_file):
    doc = docx.Document(docx_file)
    # st.header(doc)
    allText = []
    for para in doc.paragraphs:
        allText.append(para.text)
        # st.write(allText)
    text = ' '.join(allText)
    return text


def get_text_chunks(file_text,file_name):
    text_splitter = CharacterTextSplitter(
        separator = '\n',
        chunk_size = 80,
        chunk_overlap = 30,
        length_function = len
    )
    chunks = text_splitter.split_text(file_text)
    # st.write(chunks)
    doc_list = []
    for chunk in chunks:
        refrence = {'source':file_name}
        doc_string = Document(page_content = chunk , metadata = refrence)
        doc_list.append(doc_string) 
    return doc_list

# def get_vectorstore(text_chunks):
#     # Create FAISS index
#     embeddings_array = []
#     for chunk in text_chunks:
#         embeddings_array.append(embeddings.embed_documents([chunk.page_content])[0])  # Convert each chunk to its embedding
#     embeddings_array = np.array(embeddings_array).astype('float32')

#     # Initialize FAISS index
#     dimension = embeddings_array.shape[1]
#     index = faiss.IndexFlatL2(dimension)  # L2 distance metric
#     index.add(embeddings_array)  # Add embeddings to the index

#     # Return FAISS as vectorstore
#     return FAISS(embedding_function=embeddings.embed_documents, index=index)

def get_vector_store(text_chunks):
    embeddings_list = []
    embeddings = HuggingFaceEmbeddings(model_name = 'sentence-transformers/all-mpnet-base-v2')

    for chunk in text_chunks:
        embeddings_list.append(embeddings.embed_documents([chunk.page_content])[0])
        return embeddings_list




if __name__ == '__main__':
    main()