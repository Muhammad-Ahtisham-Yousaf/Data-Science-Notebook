import streamlit as st 
# from qdrant_client import QdrantClient
import os

from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
from datetime import datetime
import random
import string
from streamlit_chat import message
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document
from langchain.vectorstores import Qdrant
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI



embedding_model = HuggingFaceEmbeddings(model_name = "sentence-transformers/all-mpnet-base-v2" )



openapi_key = st.secrets["OPENAI_API_KEY"]
qdrant_url = st.secrets['QDRANT_URL']
qdrant_key = st.secrets['QDRANT_API_KEY']

def main():
    load_dotenv()
    st.set_page_config("MY Q/A ChatBot")
    st.title("Q/A Chain")

    if 'conversation' not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = None
    if "processComplete" not in st.session_state:
        st.session_state.processComplete = None
    
    with st.sidebar:
        uploaded_files = st.file_uploader("Upload Your File Here:",type=['pdf','docx'],accept_multiple_files = True)
        
        openai_api_key = openapi_key
        process = st.button("Start Process")
        
    if process:
        if not uploaded_files:
            st.warning("Please upload file first!")
            st.stop()
        if not openai_api_key:
            st.warning("Please provide your API key first!")
            st.stop()

        text_chunks_list = []
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_text = get_files_text(uploaded_file)
            st.write("File Text Extracted...")
            text_chunks = get_text_chunks(file_text,file_name)
            st.write("Text Chunks Created..")
            text_chunks_list.extend(text_chunks)

        curr_date = str(datetime.now())  
        my_collection_name = "".join(random.choices(string.ascii_letters,k=4)) + curr_date.split(".")[0].replace(":","-").replace(" ","T")# random letters + formated date
        vectorstore = get_vector_store(text_chunks_list,my_collection_name)
        st.write("Vector Store Created..")
        num_chunks = 4
        st.session_state.conversation =  get_QA_chain(vectorstore,num_chunks)
        st.session_state.processComplete = True

    if st.session_state.processComplete == True:
        user_question = st.chat_input("Ask Questions about your files")    
        if user_question:
           handel_user_input(user_question)







def get_files_text(uploaded_file):
    text = ""
    file_name = os.path.splitext(uploaded_file.name)
    file_extension = file_name[1]
    if file_extension == '.pdf':
        text+=get_pdf_text(uploaded_file)
    elif file_extension == 'docx':
        text+=get_docx_text(uploaded_file)
    else:
        pass
    return text


def get_pdf_text(uploaded_file):
    pdf_reader = PdfReader(uploaded_file)
    text = ""
    for page in pdf_reader.pages:
        text+= page.extract_text()
    return text

def get_docx_text(uploaded_file):
    doc = docx.Document(uploaded_file)
    allText = []
    for docparagraph in doc.paragraphs:
        allText.append(docparagraph.text) 
    text = ' '.join(allText)
    return text
    

def get_text_chunks(file_text,file_name):
    text_splitter = CharacterTextSplitter(separator = "\n",chunk_size = 280,chunk_overlap = 50,length_function = len)
    text_chunks = text_splitter.split_text(file_text)
    doc_list = []
    for text_chunk in text_chunks:
        # st.write(text_chunk.page_content)
        file_source = {"source": file_name}
        doc_in_str = Document(page_content = text_chunk,metadata = file_source)
        doc_list.append(doc_in_str)
    return doc_list


def get_vector_store(text_chunks,my_collection_name):
    try:
        vector_database = Qdrant.from_documents(
            documents=text_chunks,
            embedding= embedding_model,
            url = qdrant_url,
            prefer_grpc = True,
            api_key = qdrant_key,
            collection_name = my_collection_name
        )
    except Exception as e:
        st.write(f"Error: {e}")
    return vector_database

def get_QA_chain(vectorstore,num_chunks):
    QA = RetrievalQA.from_chain_type(
        llm = ChatOpenAI(model_name = 'gpt-3.5-turbo'),#default
        chain_type= "stuff",
        retriever = vectorstore.as_retriever(search_type = "similarity",search_kwargs = {"k":num_chunks}),
        return_source_documents = True
    )
    return QA 


def handel_user_input(user_question):
    try:
        with st.spinner("Thinking..."):
            result = st.session_state.conversation({"query": user_question})
            bot_response = result['result']
            source = result['source_document'][0].metadata['source']
            st.session_state.chat_history.append(user_question)
            st.session_state.chat_history.append(f"{bot_response} \n Source Document:{source}")

            response_container = st.container()

            with response_container:
                for i ,messages in enumerate(st.session_state.chat_history):
                    if i % 2 == 0 :
                        message(messages,key = str(i),is_user = True)
                    else:
                        message(messages , key= str(i))
    except:
        st.info("You Have eccedded your current qouta.")

if __name__ == '__main__':
    main()